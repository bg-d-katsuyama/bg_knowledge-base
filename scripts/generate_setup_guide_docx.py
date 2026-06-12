"""久保田様向けセットアップ手順書（Word）を生成するスクリプト。

実行方法:
    uv run python scripts/generate_setup_guide_docx.py

出力先:
    output/BGナレッジベース_セットアップ手順書_久保田様向け_v1.1.docx

v1.1 (2026-06-12): 久保田様の PC が Mac と判明したため macOS 前提に全面改訂。
    勝山様記入欄（受け渡し情報シート）と勝山様向け準備手順（送付前に削除する章）を追加。

注意: 本スクリプトは引き継ぎドキュメント生成専用の一回限りのユーティリティ。
外部 API は一切呼ばない（コストゼロ）。
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

JP_FONT = "Yu Gothic"
CODE_FONT = "Consolas"
FILL = "＿" * 24  # 勝山様記入欄の下線

OUTPUT_PATH = Path(__file__).resolve().parent.parent / "output" / (
    "BGナレッジベース_セットアップ手順書_久保田様向け_v1.1.docx"
)


def set_style_font(doc: Document, style_name: str, size: Pt | None = None) -> None:
    """指定スタイルの欧文・和文フォントを揃える。"""
    style = doc.styles[style_name]
    style.font.name = JP_FONT
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), JP_FONT)
    if size is not None:
        style.font.size = size


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def step(doc: Document, number: str, text: str) -> None:
    """「手順 N.」形式の段落(番号は手動で振る)。"""
    para = doc.add_paragraph()
    run = para.add_run(f"{number} ")
    run.bold = True
    para.add_run(text)


def code(doc: Document, text: str) -> None:
    """グレー背景のコマンド表示ボックス。"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    rfonts = run._element.get_or_add_rPr()
    el = OxmlElement("w:rFonts")
    el.set(qn("w:ascii"), CODE_FONT)
    el.set(qn("w:hAnsi"), CODE_FONT)
    el.set(qn("w:eastAsia"), "MS Gothic")
    rfonts.append(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    para._p.get_or_add_pPr().append(shd)


def prompt_box(doc: Document, text: str) -> None:
    """Claude Code への依頼文(プロンプト)を示す薄青背景ボックス。"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "E8F0FE")
    para._p.get_or_add_pPr().append(shd)


def note(doc: Document, text: str, label: str = "ポイント", fill: str = "FFF6DD") -> None:
    """注意書きボックス。"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(f"【{label}】 ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
    para.add_run(text)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    para._p.get_or_add_pPr().append(shd)


def check(doc: Document, text: str) -> None:
    """チェックリスト項目。"""
    doc.add_paragraph(f"□　{text}")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            t.rows[r].cells[c].text = value
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    for name in ("Normal", "List Bullet"):
        set_style_font(doc, name, Pt(10.5))
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        set_style_font(doc, name)

    # ---- 表紙 ----
    doc.add_heading("BGナレッジベース 運用セットアップ手順書", level=0)
    p(doc, "（久保田様向け・移行期間運用版・Mac対応）", bold=True)
    doc.add_paragraph()
    table(
        doc,
        ["項目", "内容"],
        [
            ["版", "v1.1（Mac対応・受け渡し情報シート追加）"],
            ["作成日", "2026年6月12日"],
            ["作成", "勝山（Claude Code により作成）"],
            ["対象読者", "久保田様（運用責任者）"],
            ["前提", "macOS のパソコン（MacBook 等）"],
        ],
    )

    # ---- 勝山様向け（削除前提の特出し章） ----
    h1(doc, "【勝山様向け】送付前の準備（この章は完了後に削除してから共有）")
    note(
        doc,
        "この章は勝山様の作業用です。以下をすべて完了し、第4章の「受け渡し情報シート」への記入を済ませたら、"
        "この章を丸ごと削除（章タイトルからA-7の末尾まで選択して Delete）してから久保田様へ送付してください。",
        label="この章の扱い",
        fill="FDECEC",
    )
    h2(doc, "A-1. 事前確認（久保田様への確認結果を控える）")
    p(doc, f"・久保田様の GitHub アカウント名：{FILL}")
    p(doc, "・Claude の有料プラン（Pro 以上）加入：□ 確認済み")
    p(doc, "・PC：Mac（確認済み・本書は Mac 前提）")
    h2(doc, "A-2. GitHub コラボレータ招待")
    step(doc, "1.", "ブラウザで GitHub に bg-d-katsuyama としてログインします。")
    step(
        doc,
        "2.",
        "リポジトリ bg_knowledge-base の Settings → Collaborators → Add people で、"
        "A-1 で確認した久保田様のアカウント名を検索し、Role は Write で招待します。",
    )
    step(doc, "3.", "久保田様に招待メールの承諾（Accept invitation）を依頼します。")
    h2(doc, "A-3. 久保田様専用 Notion インテグレーションの作成（読み取り専用）")
    step(doc, "1.", "ブラウザで https://www.notion.so/profile/integrations を開きます（BG ワークスペースの管理者でログイン）。")
    step(doc, "2.", "「新しいインテグレーション」を作成します。名前の例：「BG-KB-kubota」。関連ワークスペースは BG を選択します。")
    step(
        doc,
        "3.",
        "「機能」（Capabilities）で「コンテンツを読み取る」のみを有効にし、"
        "「コンテンツを更新する」「コンテンツを挿入する」は必ずオフにします。"
        "これにより、誤操作でも Notion を書き換えられないことがトークンのレベルで保証されます。",
    )
    step(doc, "4.", "作成後に表示される「内部インテグレーションシークレット」（ntn_ で始まる文字列）がトークンです。コピーして控えます。")
    step(
        doc,
        "5.",
        "対象ページ・DB にこのインテグレーションを接続します。各ページ右上の「…」→「接続」→「BG-KB-kubota」を選択。"
        "対象：KB 5 DB（ナレッジエントリ／人／企業・団体／プロジェクト／タグ）、MTG Minutes DB、"
        "腐植性堆肥製造マニュアル v1.3 のページ。"
        "（親ページに接続すると配下のページにも適用されます。5 DB が同じ親ページ配下なら親への接続1回で済みます）",
    )
    h2(doc, "A-4. 受け渡し情報シート（第4章）への記入")
    step(doc, "1.", "勝山様の .env から NOTION_DB_〜 の5つの ID を第4章のシートに転記します。")
    step(
        doc,
        "2.",
        "Notion トークンは原則このシートに書かず、別経路（パスワード管理ツールの共有、口頭＋短時間のメッセージ等）で渡します。"
        "本書に書く場合は、本書自体をトークンと同じ機密度で扱ってください（メール添付は避け、共有後に久保田様側で保管場所を限定）。",
    )
    h2(doc, "A-5. 引き継ぎ zip の送付")
    step(
        doc,
        "1.",
        "output/handover_kubota_20260612.zip（約23MB）を久保田様へ渡します。"
        "メール添付にはサイズが大きいため、Google Drive での共有を推奨します。"
        "渡した場所を第4章のシートに記入してください。",
    )
    h2(doc, "A-6. 送付前チェック")
    check(doc, "A-1〜A-5 がすべて完了している")
    check(doc, "第4章の受け渡し情報シートに記入した")
    check(doc, "トークンの渡し方を決めて実行した")
    h2(doc, "A-7. この章を削除して保存し、久保田様へ送付")
    p(doc, "以上で勝山様の準備は完了です。この章を削除し、上書き保存してから送付してください。")

    # ---- 1. この手順書について ----
    h1(doc, "1. この手順書について")
    p(
        doc,
        "この手順書は、BGナレッジベース（議事録から知見を抽出してマスターデータを作る仕組み）を、"
        "久保田様ご自身の Mac で操作できるようにするためのものです。"
        "プログラミングの知識は不要です。書いてある通りに、上から順番に進めてください。",
    )
    bullet(doc, "初期セットアップ（第5章〜第12章）の所要時間：約60〜90分")
    bullet(doc, "日常の運用（第13章）の所要時間：1回あたり10〜30分程度")
    note(
        doc,
        "途中でエラーが出たり、画面の表示が手順書と違ったりしたら、無理に進めず勝山までご連絡ください。"
        "移行期間中は勝山がいつでもサポートします。",
        label="大事なこと",
    )

    # ---- 2. 全体像 ----
    h1(doc, "2. システムの全体像")
    p(
        doc,
        "このシステムの目的は、Google Meet の議事録（Word ファイル）から土壌R&Dの知見を抽出し、"
        "マスターデータ（Excel ファイル）にまとめることです。日常運用の流れは次の4ステップです。",
    )
    step(doc, "①", "Google Drive から新しい議事録ファイル（.docx）をダウンロードする")
    step(doc, "②", "Claude Code（AIアシスタント）に知見の抽出を頼む")
    step(doc, "③", "出来上がった Excel ファイル（output フォルダ）を確認する")
    step(doc, "④", "内容をマスターデータのスプレッドシートに転記する")
    doc.add_paragraph()
    h2(doc, "Claude Code とは")
    p(
        doc,
        "Claude Code は、Anthropic 社の AI「Claude」に、パソコン上のファイル操作を日本語でお願いできるツールです。"
        "「ターミナル」という文字だけの画面で動きますが、コマンドを覚える必要はありません。"
        "「〇〇のファイルから知見を抽出して」と日本語で話しかけるだけで作業してくれます。",
    )
    note(
        doc,
        "このシステムは外部の API 課金（従量課金）を使いません。Claude の月額プラン（Pro など）の範囲内で動作します。"
        "また、Notion のデータは「読み取り専用」で、既存のページを書き換えることはありません。",
    )

    # ---- 3. 移行期間の体制 ----
    h1(doc, "3. 移行期間の体制とルール")
    p(
        doc,
        "いきなり完全移管はせず、当面は「移行期間」として、勝山と久保田様の両方が操作できる体制にします。"
        "二人で同じものを扱うため、次のルールを守ってください。",
    )
    table(
        doc,
        ["ルール", "理由"],
        [
            [
                "作業を始める前に、Slack 等で「今からKB作業をします」と一声かける",
                "二人が同時に作業すると、結果がぶつかることがあるため",
            ],
            [
                "作業の最初に Claude Code へ「git pull して最新の状態にしてください」と頼む",
                "もう一人が行った変更を取り込んでから作業を始めるため",
            ],
            [
                "作業の最後に「今回の作業内容を docs/decision_log.md に記録し、コミットしてプッシュしてください」と頼む",
                "作業の記録を残し、もう一人に変更を共有するため",
            ],
            [
                "Notion のトークン（後述）は各自専用のものを使い、共有しない",
                "問題があったときに、どちらの操作か切り分けられるようにするため",
            ],
            [
                "判断に迷ったら手を止めて勝山に連絡する",
                "移行期間中は勝山がサポートに入れるため",
            ],
        ],
    )
    p(
        doc,
        "移行期間が終わって完全移管するときに、勝山側のアクセス権の整理（GitHub・Notion トークンの無効化）を行います。"
        "それまでは勝山側の環境もそのまま残します。",
    )

    # ---- 4. 事前準備＋受け渡し情報シート ----
    h1(doc, "4. 事前準備チェックリストと受け渡し情報シート")
    h2(doc, "4-1. ご自身で用意するもの")
    check(doc, "macOS のパソコン（インターネットに接続できること）")
    check(doc, "GitHub のアカウント（無料）。https://github.com/ で「Sign up」から作成できます")
    check(
        doc,
        "Claude のアカウント（有料プラン）。https://claude.ai/ で登録し、Pro 以上のプランに加入してください"
        "（Claude Code の利用に必要です。プランは会社のご判断で構いません）",
    )
    check(doc, "Notion に BG ワークスペースのメンバーとしてログインできること")
    h2(doc, "4-2. 勝山から届くもの")
    table(
        doc,
        ["届くもの", "何に使うか"],
        [
            [
                "GitHub リポジトリへの招待メール",
                "届いたら「Accept invitation」を押してください（第10章で使います）",
            ],
            [
                "Notion トークン（久保田様専用・読み取り専用）",
                "Notion のデータを読み取るための「鍵」。パスワードと同じ扱いで管理してください",
            ],
            [
                "引き継ぎファイル一式（zip）",
                "過去の抽出結果と議事録ファイル。第11章で配置します",
            ],
        ],
    )
    h2(doc, "4-3. 受け渡し情報シート（勝山が記入してお渡しします）")
    p(doc, "セットアップの第11章（.env の作成）でこの表の値を使います。")
    table(
        doc,
        ["設定名", "値"],
        [
            ["リポジトリの場所（記入済み）", "https://github.com/bg-d-katsuyama/bg_knowledge-base.git"],
            ["NOTION_API_TOKEN", "※原則ここには記載せず、別経路でお渡しします"],
            ["NOTION_DB_KNOWLEDGE_ENTRY（ナレッジエントリDB）", FILL],
            ["NOTION_DB_PEOPLE（人DB）", FILL],
            ["NOTION_DB_ORGANIZATION（企業・団体DB）", FILL],
            ["NOTION_DB_PROJECT（プロジェクトDB）", FILL],
            ["NOTION_DB_TAG（タグDB）", FILL],
            ["引き継ぎ zip の置き場所（共有方法）", FILL],
        ],
    )
    note(
        doc,
        "この表に記入された値のうち DB の ID は秘密情報ではありませんが、トークンはパスワードと同じです。"
        "トークンが書かれた状態の本書やメモを、メール添付や共有フォルダに置きっぱなしにしないでください。",
        label="重要",
        fill="FDECEC",
    )

    # ---- 5. ターミナル ----
    h1(doc, "5. ステップ1：ターミナルの開き方")
    p(
        doc,
        "この後の手順では「ターミナル」という、文字でパソコンに指示を出す画面を使います。"
        "開き方を覚えてください（今後何度も使います）。",
    )
    step(doc, "1.", "キーボードで command（⌘）キーを押しながらスペースキーを押します（Spotlight 検索が開きます）。")
    step(doc, "2.", "「ターミナル」と入力し、Enter キーを押します。")
    step(doc, "3.", "白か黒の画面が開き、「（お名前）@〜 %」のような文字が表示されれば成功です。")
    note(
        doc,
        "この手順書でグレーの枠に書かれている「コマンド」は、この画面にコピーして貼り付け（⌘V）、"
        "Enter キーを押して実行します。1行ずつ実行してください。",
    )

    # ---- 6. Git ----
    h1(doc, "6. ステップ2：Git のインストール")
    p(
        doc,
        "Git（ギット）は、プログラム一式を取得・更新するための道具です。Mac には半自動で入れられます。",
    )
    step(doc, "1.", "ターミナルで次のコマンドを実行します。")
    code(doc, "git --version")
    step(
        doc,
        "2.",
        "「git version 2.〜」のように数字が表示されたら、すでにインストール済みです。第7章へ進んでください。",
    )
    step(
        doc,
        "3.",
        "「コマンドライン・デベロッパツールをインストールしますか？」というウインドウが出た場合は、"
        "「インストール」を押し、利用規約に同意して完了を待ちます（数分〜数十分かかることがあります）。",
    )
    step(doc, "4.", "完了したらターミナルをいったん閉じて開き直し、もう一度 git --version を実行して数字が出ることを確認します。")

    # ---- 7. uv ----
    h1(doc, "7. ステップ3：uv のインストール")
    p(
        doc,
        "uv（ユーブイ）は、このシステムが使う Python（プログラムの実行環境）を自動で整えてくれる道具です。",
    )
    step(doc, "1.", "ターミナルで次のコマンドを実行します。")
    code(doc, "curl -LsSf https://astral.sh/uv/install.sh | sh")
    step(doc, "2.", "完了したらターミナルを閉じて開き直し、次のコマンドで確認します。")
    code(doc, "uv --version")
    p(doc, "「uv 0.〜」のように数字が表示されれば成功です。")

    # ---- 8. Claude Code ----
    h1(doc, "8. ステップ4：Claude Code のインストール")
    step(doc, "1.", "ターミナルで次のコマンドを実行します。")
    code(doc, "curl -fsSL https://claude.ai/install.sh | bash")
    step(doc, "2.", "完了したらターミナルを閉じて開き直し、次のコマンドで確認します。")
    code(doc, "claude --version")
    p(doc, "数字が表示されれば成功です。ログインはプロジェクト取得後の第12章で行います。")

    # ---- 9. GitHub CLI ----
    h1(doc, "9. ステップ5：GitHub へのログイン設定（GitHub CLI）")
    p(
        doc,
        "作業記録を GitHub に送る（プッシュする）ために、一度だけログイン設定をします。"
        "「GitHub CLI」という公式の道具を使うと、ブラウザでログインするだけで設定が終わります。",
    )
    step(
        doc,
        "1.",
        "ブラウザで https://github.com/cli/cli/releases/latest を開き、ページ下部の Assets から "
        "「gh_〜_macOS_universal.pkg」をダウンロードして、ダブルクリックでインストールします（設定はすべてそのまま「続ける」で構いません）。",
    )
    step(doc, "2.", "ターミナルを開き直し、次のコマンドを実行します。")
    code(doc, "gh auth login")
    step(doc, "3.", "いくつか質問されます。次の表の通りに答えてください（矢印キーで選んで Enter）。")
    table(
        doc,
        ["質問（英語）", "選ぶ答え"],
        [
            ["Where do you use GitHub?", "GitHub.com"],
            ["What is your preferred protocol for Git operations?", "HTTPS"],
            ["Authenticate Git with your GitHub credentials?", "Y（Enter）"],
            ["How would you like to authenticate GitHub CLI?", "Login with a web browser"],
        ],
    )
    step(
        doc,
        "4.",
        "画面に8桁のコード（XXXX-XXXX）が表示されます。控えてから Enter を押すとブラウザが開くので、"
        "ご自身の GitHub アカウントでログインし、コードを入力して「Authorize」を押します。",
    )
    step(doc, "5.", "ターミナルに「Logged in as（アカウント名）」と表示されれば成功です。")

    # ---- 10. クローン ----
    h1(doc, "10. ステップ6：プロジェクトの取得（クローン）")
    p(
        doc,
        "GitHub に保管されているプログラム一式を、ご自身の Mac にコピーします。これを「クローン」と呼びます。"
        "事前に GitHub の招待メール（4-2参照）を承諾（Accept invitation）しておいてください。",
    )
    step(doc, "1.", "ターミナルで次の2つのコマンドを順に実行します（書類フォルダに移動してから取得します）。")
    code(doc, "cd ~/Documents")
    code(doc, "git clone https://github.com/bg-d-katsuyama/bg_knowledge-base.git")
    step(doc, "2.", "エラーなく完了したら、次のコマンドでフォルダに移動します。")
    code(doc, "cd bg_knowledge-base")
    note(
        doc,
        "今後、Claude Code を使うときは毎回、ターミナルを開いて「cd ~/Documents/bg_knowledge-base」で"
        "このフォルダに移動してから始めます。この1行は覚えておいてください（コピーして使えば大丈夫です）。",
    )

    # ---- 11. 初期設定 ----
    h1(doc, "11. ステップ7：プロジェクトの初期設定")
    h2(doc, "11-1. 必要な部品のインストール")
    p(doc, "プロジェクトのフォルダ（bg_knowledge-base）にいる状態で、次のコマンドを実行します。")
    code(doc, "uv sync")
    p(
        doc,
        "Python 本体や必要な部品が自動でダウンロードされます。初回は数分かかります。"
        "エラーらしき赤い文字が出なければ成功です。",
    )
    h2(doc, "11-2. 設定ファイル（.env）の作成")
    p(
        doc,
        ".env（ドットエンブ）は、Notion トークンなどの秘密情報を書いておくファイルです。"
        "第4章の「受け渡し情報シート」と、別途受け取った Notion トークンを手元に用意してから進めてください。",
    )
    step(doc, "1.", "次のコマンドで、ひな形をコピーして .env を作ります。")
    code(doc, "cp .env.example .env")
    step(doc, "2.", "次のコマンドでテキストエディット（Mac のメモ帳）が開きます。")
    code(doc, "open -e .env")
    step(
        doc,
        "3.",
        "受け渡し情報シートの値を、対応する行の「=」の右側に貼り付けます。対象は次の6行です。"
        "それ以外の行は変更せず、空欄のままで構いません。",
    )
    table(
        doc,
        ["行（設定名）", "貼り付ける値"],
        [
            ["NOTION_API_TOKEN", "別途受け取った久保田様専用トークン"],
            ["NOTION_DB_KNOWLEDGE_ENTRY", "シートの「ナレッジエントリDB」の値"],
            ["NOTION_DB_PEOPLE", "シートの「人DB」の値"],
            ["NOTION_DB_ORGANIZATION", "シートの「企業・団体DB」の値"],
            ["NOTION_DB_PROJECT", "シートの「プロジェクトDB」の値"],
            ["NOTION_DB_TAG", "シートの「タグDB」の値"],
        ],
    )
    step(doc, "4.", "⌘S で保存して、テキストエディットを閉じます。")
    note(
        doc,
        ".env はパスワードと同じです。メールに添付したり、Slack に貼り付けたりしないでください。"
        "値の前後に余計なスペースが入ると認証エラーになるので注意してください。",
        label="重要",
        fill="FDECEC",
    )
    h2(doc, "11-3. 引き継ぎファイルの配置")
    p(
        doc,
        "勝山から受け取った「引き継ぎファイル一式」（zip）をダブルクリックで展開し、"
        "中のファイルを次の場所に置いてください。Finder（フォルダ画面）でのコピーで構いません。"
        "プロジェクトのフォルダは「書類（Documents）→ bg_knowledge-base」にあります。",
    )
    table(
        doc,
        ["zip の中のフォルダ", "コピー先"],
        [
            [
                "logs フォルダの中身（insights_drive.json など）",
                "bg_knowledge-base の logs フォルダの中（フォルダがなければ新規作成）",
            ],
            [
                "data/drive_input フォルダの中身（議事録 docx 一式）",
                "bg_knowledge-base の data → drive_input フォルダの中",
            ],
            [
                "output フォルダの中身（過去の Excel ファイル）",
                "bg_knowledge-base の output フォルダの中（参考用）",
            ],
        ],
    )

    # ---- 12. 動作確認 ----
    h1(doc, "12. ステップ8：Claude Code の起動と動作確認")
    h2(doc, "12-1. 起動とログイン")
    step(doc, "1.", "プロジェクトのフォルダにいる状態で、次のコマンドを実行します。")
    code(doc, "claude")
    step(
        doc,
        "2.",
        "初回はログインを求められます。「Claude account with subscription」を選ぶとブラウザが開くので、"
        "ご自身の Claude アカウント（Pro 以上）でログインし、許可（Authorize）してください。",
    )
    step(
        doc,
        "3.",
        "画面の配色（テーマ）の選択などを聞かれたら、そのまま Enter で進めて構いません。"
        "「Do you trust the files in this folder?（このフォルダを信頼しますか）」と聞かれたら「Yes, proceed」を選びます。",
    )
    step(doc, "4.", "入力欄が表示されたら起動成功です。ここに日本語で話しかけて作業を進めます。")
    h2(doc, "12-2. Claude Code の「許可確認」について")
    p(
        doc,
        "Claude Code は、ファイルを変更したりコマンドを実行したりする前に「許可しますか？」と確認してきます。"
        "次の方針で答えてください。",
    )
    table(
        doc,
        ["確認の内容", "答え方"],
        [
            ["ファイルの読み取り、output フォルダへのファイル作成", "許可してOK（Yes）"],
            ["docs/decision_log.md への記録、git のコミット・プッシュ", "許可してOK（Yes）"],
            ["Notion への書き込み・更新・削除に見えるもの", "許可しない（No）→ 勝山に連絡"],
            ["内容がよく分からないもの", "許可しない（No）→ 勝山に連絡"],
        ],
    )
    h2(doc, "12-3. 動作確認テスト")
    step(doc, "1.", "まず、次の文をそのまま入力欄に貼り付けて Enter を押してください。")
    prompt_box(
        doc,
        "CLAUDE.md と docs/decision_log.md を読んで、このプロジェクトの目的と現在の状況を、"
        "専門用語を使わずに5行以内で説明してください。",
    )
    p(doc, "プロジェクトの説明が日本語で返ってくれば、Claude Code は正しく動いています。")
    step(doc, "2.", "次に、ファイルが読めることを確認します。")
    prompt_box(
        doc,
        "data/drive_input フォルダの中の docx ファイルを1つ選んで、内容を3行で要約してください。"
        "Notion への書き込みは行わないでください。",
    )
    p(doc, "議事録の要約が返ってくれば、セットアップはすべて完了です。お疲れさまでした。")
    note(
        doc,
        "Claude Code を終了するには、入力欄に /exit と入力するか、ターミナルの画面を閉じます。"
        "会話をリセットして新しく始めたいときは /clear と入力します。",
    )

    # ---- 13. 日常運用 ----
    h1(doc, "13. 日常の運用手順（新しい議事録から知見を抽出する）")
    p(doc, "新しい議事録が増えたときの定常作業です。毎回この順番で行ってください。")
    step(doc, "1.", "Slack 等で「今からKB作業をします」と一声かけます（移行期間ルール）。")
    step(
        doc,
        "2.",
        "Google Drive から新しい議事録ファイル（.docx）をダウンロードし、"
        "Finder で「書類 → bg_knowledge-base → data → drive_input」フォルダにコピーします。",
    )
    step(doc, "3.", "ターミナルを開き、次の2つを実行して Claude Code を起動します。")
    code(doc, "cd ~/Documents/bg_knowledge-base")
    code(doc, "claude")
    step(doc, "4.", "最初に、次の文を貼り付けて最新の状態に更新します。")
    prompt_box(doc, "git pull して、プロジェクトを最新の状態にしてください。")
    step(doc, "5.", "続けて、次の文を貼り付けます（ファイル名の部分はご自身で書き換えてください）。")
    prompt_box(
        doc,
        "data/drive_input に新しい議事録ファイル「（ここにファイル名）.docx」を追加しました。"
        "この内容から土壌R&Dの知見を抽出し、これまでの知見（logs/insights_drive.json）と統合した"
        "マスターデータの xlsx を output フォルダに作成してください。"
        "Notion への書き込みは行わないでください。",
    )
    step(
        doc,
        "6.",
        "完了したら、output フォルダにできた新しい Excel ファイルを開き、抽出された知見の内容を確認します。"
        "おかしな内容があれば、Claude Code に日本語で修正を頼めます（例：「3行目の知見は議事録の文脈と違うので削除して」）。",
    )
    step(doc, "7.", "問題なければ、マスターデータのスプレッドシートに転記します。")
    step(doc, "8.", "最後に、次の文を貼り付けて作業記録を残します。")
    prompt_box(
        doc,
        "今回の作業内容を docs/decision_log.md に記録して、変更をコミットしてプッシュしてください。",
    )

    # ---- 14. プロンプト集 ----
    h1(doc, "14. よく使う依頼文（プロンプト）集")
    p(doc, "Claude Code への依頼は自由な日本語で構いませんが、定型文として以下が便利です。")
    table(
        doc,
        ["やりたいこと", "依頼文の例"],
        [
            [
                "最新状態への更新（毎回最初に）",
                "git pull して、プロジェクトを最新の状態にしてください。",
            ],
            [
                "新しい議事録からの知見抽出",
                "第13章の手順5の文をご利用ください。",
            ],
            [
                "抽出結果の手直し",
                "output の最新の xlsx の〇行目について、（修正内容）に修正して再生成してください。",
            ],
            [
                "現状の確認",
                "docs/decision_log.md を読んで、いま何が完了していて、次に何をすべきか教えてください。",
            ],
            [
                "作業記録と共有（毎回最後に）",
                "今回の作業内容を docs/decision_log.md に記録して、変更をコミットしてプッシュしてください。",
            ],
            [
                "分からないことの質問",
                "（なんでも日本語で質問できます。例：このプロジェクトでの「タグDB」の役割を教えて）",
            ],
        ],
    )

    # ---- 15. やってはいけないこと ----
    h1(doc, "15. やってはいけないこと")
    bullet(
        doc,
        "Notion の既存ページ・データベースを Claude Code 経由で編集・削除すること。"
        "このシステムは Notion を「読み取り専用」で使う設計です。書き込みを求められても許可しないでください"
        "（久保田様のトークンは読み取り専用で発行されているため、誤って許可しても書き込みはできない仕組みです）。",
    )
    bullet(
        doc,
        ".env ファイルや Notion トークンを、メール・Slack・チャットに貼り付けて共有すること（パスワードと同じ扱いです）。",
    )
    bullet(
        doc,
        "docs/architecture.md（設計書の正本）、infra フォルダ、.github/workflows/deploy.yml の変更。"
        "変更が必要な場合は勝山に相談してください。",
    )
    bullet(
        doc,
        "Claude Code の許可確認で、内容が理解できない操作を許可すること。迷ったら No を選び、勝山に確認してください。",
    )
    bullet(
        doc,
        "エラーが出た状態で操作を繰り返すこと。同じエラーが2回出たら手を止めて勝山に連絡してください。",
    )

    # ---- 16. トラブルシューティング ----
    h1(doc, "16. 困ったときは（トラブルシューティング）")
    table(
        doc,
        ["症状", "対処"],
        [
            [
                "「command not found: git」「command not found: claude」等と表示される",
                "ターミナルをいったん閉じて開き直してください。それでも直らなければ該当ステップのインストールをやり直してください。",
            ],
            [
                "ダウンロードした pkg が「開発元を確認できないため開けません」と言われる",
                "ファイルを右クリック（または control キーを押しながらクリック）→「開く」を選ぶと進めます。",
            ],
            [
                "uv sync でエラーが出る",
                "インターネット接続を確認して再実行。直らなければエラー画面のスクリーンショットを撮って勝山へ。",
            ],
            [
                "Notion 関連で「unauthorized」「401」等のエラーが出る",
                ".env のトークンが正しいか（前後に空白がないか）を確認。直らなければ勝山へ（トークンの再発行で解決することがあります）。",
            ],
            [
                "git pull で「conflict（競合）」と表示される",
                "そのまま触らず勝山に連絡してください（二人の変更がぶつかった状態です）。",
            ],
            [
                "Claude Code の応答が止まった・様子がおかしい",
                "esc キーで中断できます。/clear で新しい会話を始めるか、ターミナルを閉じて起動し直してください。",
            ],
            [
                "ログインのブラウザ画面が開かない",
                "ターミナルに表示される URL を手動でブラウザに貼り付けて開いてください。",
            ],
        ],
    )
    p(doc, "上記で解決しない場合の連絡先：勝山（GitHub: bg-d-katsuyama）。移行期間中は遠慮なくご連絡ください。", bold=True)

    # ---- 17. 用語集 ----
    h1(doc, "17. 用語集")
    table(
        doc,
        ["用語", "意味"],
        [
            ["ターミナル", "文字でパソコンに指示を出す画面。⌘＋スペース →「ターミナル」で開く"],
            ["コマンド", "ターミナルに入力する指示文。この手順書ではグレーの枠で示している"],
            ["Finder", "Mac のフォルダ・ファイルを見る画面（Windows のエクスプローラーに相当）"],
            ["リポジトリ", "プログラムや文書の一式を保管する場所。GitHub 上に正本がある"],
            ["クローン", "GitHub のリポジトリを自分のパソコンにコピーすること"],
            ["git pull / プッシュ", "pull は最新を取り込むこと、プッシュ（push）は自分の変更を GitHub に送ること"],
            ["コミット", "変更内容に説明を付けて記録すること"],
            ["トークン", "システムにアクセスするための「鍵」。パスワードと同じ扱いで管理する"],
            [".env", "トークンなどの秘密情報を書いておく設定ファイル。共有・コミット禁止"],
            ["uv", "Python（プログラム実行環境）を自動で整えてくれる道具"],
            ["GitHub CLI（gh）", "GitHub へのログイン設定を簡単にしてくれる公式の道具"],
            ["プロンプト", "Claude Code への依頼文のこと。日本語で自由に書いてよい"],
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
