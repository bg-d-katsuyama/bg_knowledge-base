"""data/drive_input 配下の .docx をローカルにテキストダンプする読み取り専用スクリプト.

Google API・GCP は使わない。python-docx でローカルファイルを読むだけ。
ダンプ結果を Claude Code（人手抽出）がそのまま読めるよう UTF-8 テキストで保存する。

段落に加えて表（議事録の決定事項・アクション一覧が表で書かれることが多い）も
タブ区切りで書き出す。

実行方法:
    # 1 ファイルだけ
    uv run python scripts/dump_docx.py "data/drive_input/<file>.docx"
    # フォルダ内の全 .docx をまとめて 1 ファイルに
    uv run python scripts/dump_docx.py --all [--in-dir data/drive_input] \\
        [--out logs/drive_input_dump.txt]
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject


def _extract_text(path: Path) -> str:
    """1 つの .docx から段落と表を順序通りにテキスト抽出する."""
    doc: DocumentObject = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        lines.append("[表]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("\t".join(cells))
    return "\n".join(lines)


def _dump_one(path: Path) -> str:
    body = _extract_text(path)
    return (
        f"# {path.name}\n"
        f"path: {path}\n"
        f"body_chars: {len(body)}\n"
        f"{'=' * 60}\n\n"
        f"{body}\n"
    )


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path", nargs="?", help="単一 .docx のパス")
    parser.add_argument("--all", action="store_true", help="--in-dir 内の全 .docx を処理")
    parser.add_argument("--in-dir", default="data/drive_input")
    parser.add_argument("--out", default="logs/drive_input_dump.txt")
    args = parser.parse_args()

    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    if args.all:
        in_dir = Path(args.in_dir)
        paths = sorted(in_dir.glob("*.docx"))
        if not paths:
            print(f"{in_dir} に .docx が見つかりません", file=sys.stderr)
            return 1
        chunks = [_dump_one(p) for p in paths]
        out_path = Path(args.out)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(("\n\n" + "#" * 60 + "\n\n").join(chunks), encoding="utf-8")
        print(f"対象ファイル数: {len(paths)}")
        print(f"出力: {out_path}")
        return 0

    if not args.docx_path:
        print("docx_path か --all のいずれかを指定してください", file=sys.stderr)
        return 1

    path = Path(args.docx_path)
    if not path.exists():
        print(f"ファイルが見つかりません: {path}", file=sys.stderr)
        return 1
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(_dump_one(path), encoding="utf-8")
    print(f"ファイル: {path.name}")
    print(f"出力: {out_path}")
    return 0


if __name__ == "__main__":
    return_code = main()
    sys.exit(return_code)
